"""
Genera el documento del Proyecto Final en formato .docx con el estilo exigido:
Calibri 11, interlineado sencillo, titulos en Calibri Light 16, parrafos
justificados y hojas numeradas.

Uso:  python3 documentacion/generar_documento.py
"""

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

AQUI = Path(__file__).resolve().parent
sys.path.insert(0, str(AQUI))
import contenido as C  # noqa: E402

AZUL = RGBColor(0x1F, 0x3A, 0x64)
GRIS = RGBColor(0x59, 0x59, 0x59)
SALIDA = AQUI / "Proyecto_Final_Programacion_III_Triana_Garcia.docx"


# ===========================================================================
# Utilidades de formato
# ===========================================================================
def preparar_estilos(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    pf = normal.paragraph_format
    pf.line_spacing = 1.0                     # interlineado sencillo
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for nivel, tamano, color in (("Heading 1", 16, AZUL),
                                 ("Heading 2", 14, AZUL),
                                 ("Heading 3", 12, GRIS)):
        estilo = doc.styles[nivel]
        estilo.font.name = "Calibri Light"     # titulos en Calibri Light
        estilo.font.size = Pt(tamano)
        estilo.font.bold = True
        estilo.font.color.rgb = color
        estilo.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri Light")
        estilo.paragraph_format.space_before = Pt(14 if nivel == "Heading 1" else 10)
        estilo.paragraph_format.space_after = Pt(6)
        estilo.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        estilo.paragraph_format.keep_with_next = True


def numerar_paginas(seccion):
    """Inserta el campo PAGE en el pie: 'Pagina N'."""
    pie = seccion.footer.paragraphs[0]
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pie.text = ""
    corrida = pie.add_run("Página ")
    corrida.font.size = Pt(9)
    corrida.font.color.rgb = GRIS

    campo = pie.add_run()
    campo.font.size = Pt(9)
    campo.font.color.rgb = GRIS
    inicio = OxmlElement("w:fldChar")
    inicio.set(qn("w:fldCharType"), "begin")
    instruccion = OxmlElement("w:instrText")
    instruccion.set(qn("xml:space"), "preserve")
    instruccion.text = "PAGE"
    fin = OxmlElement("w:fldChar")
    fin.set(qn("w:fldCharType"), "end")
    campo._r.append(inicio)
    campo._r.append(instruccion)
    campo._r.append(fin)


def parrafo(doc, texto, negrita=False, tamano=11, alineacion=None,
            espacio_despues=6, cursiva=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(espacio_despues)
    if alineacion is not None:
        p.alignment = alineacion
    r = p.add_run(texto)
    r.bold = negrita
    r.italic = cursiva
    r.font.size = Pt(tamano)
    if color is not None:
        r.font.color.rgb = color
    return p


def vinetas(doc, elementos, estilo="List Bullet"):
    for texto in elementos:
        p = doc.add_paragraph(texto, style=estilo)
        p.paragraph_format.space_after = Pt(3)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for r in p.runs:
            r.font.size = Pt(11)


def tabla(doc, encabezados, filas, anchos=None, tamano=9.5):
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    cabecera = t.rows[0]
    for i, texto in enumerate(encabezados):
        celda = cabecera.cells[i]
        celda.text = ""
        p = celda.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(texto)
        r.bold = True
        r.font.size = Pt(tamano)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        sombra = OxmlElement("w:shd")
        sombra.set(qn("w:fill"), "1F3A64")
        celda._tc.get_or_add_tcPr().append(sombra)

    for fila in filas:
        celdas = t.add_row().cells
        for i, valor in enumerate(fila):
            celdas[i].text = ""
            p = celdas[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(valor))
            r.font.size = Pt(tamano)

    if anchos:
        for fila in t.rows:
            for i, ancho in enumerate(anchos):
                fila.cells[i].width = Cm(ancho)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def titulo_tabla(doc, texto):
    parrafo(doc, texto, negrita=True, tamano=9.5, cursiva=True,
            espacio_despues=3, color=GRIS)


# ===========================================================================
# Secciones del documento
# ===========================================================================
def portada(doc):
    for _ in range(3):
        doc.add_paragraph()
    parrafo(doc, C.INSTITUCION, negrita=True, tamano=14,
            alineacion=WD_ALIGN_PARAGRAPH.CENTER, espacio_despues=2)
    parrafo(doc, "Tecnólogo en Desarrollo de Software", tamano=12,
            alineacion=WD_ALIGN_PARAGRAPH.CENTER, espacio_despues=40)

    parrafo(doc, "PROYECTO FINAL", negrita=True, tamano=20,
            alineacion=WD_ALIGN_PARAGRAPH.CENTER, espacio_despues=6, color=AZUL)
    parrafo(doc, C.ASIGNATURA, tamano=14,
            alineacion=WD_ALIGN_PARAGRAPH.CENTER, espacio_despues=40)

    parrafo(doc, C.PROYECTO, negrita=True, tamano=16,
            alineacion=WD_ALIGN_PARAGRAPH.CENTER, espacio_despues=8, color=AZUL)
    parrafo(doc, "Aplicación web desarrollada con metodología Agile-Scrum",
            tamano=11, cursiva=True,
            alineacion=WD_ALIGN_PARAGRAPH.CENTER, espacio_despues=60)

    datos = [("Sustentante", C.AUTOR), ("Matrícula", C.MATRICULA),
             ("Asignatura", C.ASIGNATURA), ("Facilitador", C.PROFESOR),
             ("Fecha de entrega", C.FECHA)]
    for etiqueta, valor in datos:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(f"{etiqueta}: ")
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(valor)
        r2.font.size = Pt(11)

    parrafo(doc, "Santo Domingo, República Dominicana", tamano=10,
            alineacion=WD_ALIGN_PARAGRAPH.CENTER, espacio_despues=0, color=GRIS)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


INDICE = [
    ("1.", "Introducción", 0),
    ("2.", "Estrategia de trabajo (planificación)", 0),
    ("2.1.", "Nombre del proyecto de software", 1),
    ("2.2.", "Tecnología a aplicar", 1),
    ("2.3.", "Objetivo del proyecto", 1),
    ("2.4.", "Alcance del proyecto", 1),
    ("2.5.", "Cronograma del proyecto", 1),
    ("2.6.", "Definición del primer Release", 1),
    ("2.7.", "Requerimientos funcionales y no funcionales", 1),
    ("3.", "Metodología Scrum", 0),
    ("3.1.", "Tareas a ejecutar", 1),
    ("3.2.", "Equipo de trabajo", 1),
    ("3.3.", "Herramientas de gestión", 1),
    ("3.4.", "Épicas", 1),
    ("3.5.", "Ceremonias de Scrum", 1),
    ("3.6.", "Historias de usuario", 1),
    ("4.", "Plan de pruebas", 0),
    ("4.1.", "Requerimientos y matriz de trazabilidad", 1),
    ("4.2.", "Criterios de aceptación y de rechazo", 1),
    ("4.3.", "Herramientas de pruebas y su justificación", 1),
    ("4.4.", "Cronograma de ejecución de pruebas", 1),
    ("4.5.", "Plantilla de casos de prueba", 1),
    ("4.6.", "Equipo de pruebas y responsabilidades", 1),
    ("4.7.", "Plan de automatización de pruebas", 1),
    ("4.8.", "Ejecución de las pruebas y evidencias", 1),
    ("5.", "Demostración y entregables", 0),
    ("6.", "Conclusiones", 0),
    ("7.", "Bibliografía", 0),
]


def indice(doc):
    doc.add_heading("Índice", level=1)
    for numero, titulo, nivel in INDICE:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.8 if nivel else 0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(f"{numero} {titulo}")
        r.font.size = Pt(11)
        r.bold = nivel == 0
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def introduccion(doc):
    doc.add_heading("1. Introducción", level=1)
    parrafo(doc,
            "Este documento recoge el desarrollo completo de BiblioITLA, un "
            "sistema de gestión de biblioteca construido como Proyecto Final "
            "de la asignatura Programación III. El trabajo no se limita a "
            "programar una aplicación: documenta cómo se planificó, cómo se "
            "organizó con Scrum, cómo se probó y qué resultados arrojaron esas "
            "pruebas.")
    parrafo(doc,
            "El problema de partida es concreto. Una biblioteca académica que "
            "controla sus préstamos en hojas de cálculo depende de que el "
            "personal recuerde las políticas: cuántos libros puede llevarse "
            "cada socio, cuándo vence un préstamo y cuánta mora corresponde "
            "cobrar. Ese control manual falla precisamente cuando hay más "
            "movimiento. La propuesta consiste en trasladar esas reglas al "
            "software, de modo que un préstamo que incumple una política sea "
            "sencillamente imposible de registrar.")
    parrafo(doc,
            "El proyecto se organizó en tres sprints de dos semanas siguiendo "
            "el marco Scrum, con diez historias de usuario repartidas en cuatro "
            "épicas. El primer Release cubre el circuito completo del negocio, "
            "desde el inicio de sesión hasta el cobro de la mora.")
    parrafo(doc,
            "El apartado de mayor peso es el plan de pruebas. Se construyó una "
            "batería de 99 pruebas automatizadas en tres niveles —unitarias, de "
            "integración y de extremo a extremo con navegador real— que se "
            "ejecutan automáticamente ante cada cambio del código. Ese trabajo "
            "no es decorativo: detectó dos defectos reales antes de la entrega, "
            "documentados en el apartado 4.8.")
    parrafo(doc,
            "El documento sigue el orden en que se hizo el trabajo: primero la "
            "planificación, luego la organización con Scrum, después el plan de "
            "pruebas con sus resultados y finalmente las conclusiones sobre lo "
            "que funcionó y lo que quedó pendiente.")
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def estrategia(doc):
    doc.add_heading("2. Estrategia de trabajo (planificación)", level=1)

    doc.add_heading("2.1. Nombre del proyecto de software", level=2)
    parrafo(doc, C.PROYECTO, negrita=True, tamano=12)
    parrafo(doc,
            "El nombre combina la actividad del sistema (biblioteca) con la "
            "institución para la que se plantea (ITLA). Es corto, se pronuncia "
            "con facilidad y describe sin ambigüedad de qué trata el producto.")

    doc.add_heading("2.2. Tecnología a aplicar", level=2)
    parrafo(doc,
            "La selección buscó herramientas maduras, con buena documentación "
            "y que permitieran automatizar las pruebas sin infraestructura "
            "adicional. La última columna explica por qué se eligió cada una.")
    titulo_tabla(doc, "Tabla 1. Tecnologías seleccionadas y justificación.")
    tabla(doc, ["Componente", "Tecnología", "Justificación"],
          C.TECNOLOGIA, anchos=[3.2, 3.4, 9.4])

    doc.add_heading("2.3. Objetivo del proyecto", level=2)
    parrafo(doc, "Objetivo general", negrita=True, espacio_despues=2)
    parrafo(doc, C.OBJETIVO)
    parrafo(doc, "Objetivos específicos", negrita=True, espacio_despues=2)
    vinetas(doc, C.OBJETIVOS_ESPECIFICOS)

    doc.add_heading("2.4. Alcance del proyecto", level=2)
    parrafo(doc,
            "Delimitar lo que el sistema no hace es tan importante como "
            "describir lo que hace: evita expectativas equivocadas y protege el "
            "alcance del Release durante el desarrollo.")
    parrafo(doc, "El sistema incluye:", negrita=True, espacio_despues=2)
    vinetas(doc, C.ALCANCE_INCLUYE)
    parrafo(doc, "El sistema NO incluye en este Release:", negrita=True,
            espacio_despues=2)
    vinetas(doc, C.ALCANCE_EXCLUYE)

    doc.add_heading("2.5. Cronograma del proyecto", level=2)
    titulo_tabla(doc, "Tabla 2. Cronograma de actividades, plazos y responsables.")
    tabla(doc, ["#", "Actividad", "Inicio", "Fin", "Duración", "Responsable"],
          C.CRONOGRAMA, anchos=[0.9, 6.4, 2.1, 2.1, 1.8, 2.7], tamano=9)

    doc.add_heading("2.6. Definición del primer Release", level=2)
    parrafo(doc, C.RELEASE_DESCRIPCION)
    parrafo(doc,
            "El Release 1 se compone de las diez historias de usuario del "
            "apartado 3.6, que suman 45 puntos de historia repartidos en tres "
            "sprints de dos semanas. Todas ellas están implementadas, probadas "
            "y demostradas en el video de entrega.")

    doc.add_heading("2.7. Requerimientos funcionales y no funcionales", level=2)
    parrafo(doc, "Requerimientos funcionales", negrita=True, espacio_despues=2)
    titulo_tabla(doc, "Tabla 3. Requerimientos funcionales del Release 1.")
    tabla(doc, ["ID", "Nombre", "Descripción", "Historia", "Prioridad"],
          C.REQUISITOS_FUNCIONALES, anchos=[1.3, 3.4, 8.0, 1.9, 1.4], tamano=9)

    parrafo(doc, "Requerimientos no funcionales", negrita=True, espacio_despues=2)
    titulo_tabla(doc, "Tabla 4. Requerimientos no funcionales y forma de verificarlos.")
    tabla(doc, ["ID", "Categoría", "Descripción", "Cómo se verifica"],
          C.REQUISITOS_NO_FUNCIONALES, anchos=[1.3, 2.6, 6.6, 5.5], tamano=9)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def scrum(doc):
    doc.add_heading("3. Metodología Scrum", level=1)
    parrafo(doc,
            "El proyecto se organizó con Scrum en tres sprints de dos semanas. "
            "Se eligió este marco porque el circuito de préstamos se comprendió "
            "mejor al construirlo: trabajar por incrementos permitió incorporar "
            "reglas que no estaban en el planteamiento inicial sin rehacer el "
            "trabajo ya hecho.")

    doc.add_heading("3.1. Tareas a ejecutar", level=2)
    parrafo(doc,
            "Las historias de usuario se descompusieron en veinte tareas "
            "técnicas. Una tarea es trabajo concreto para una persona, con una "
            "estimación en horas, mientras que una historia expresa valor para "
            "el usuario.")
    titulo_tabla(doc, "Tabla 5. Descomposición del trabajo en tareas técnicas.")
    tabla(doc, ["ID", "Tarea", "Rol responsable", "Sprint", "Estimación"],
          C.TAREAS, anchos=[1.2, 7.6, 3.0, 1.9, 2.0], tamano=9)

    doc.add_heading("3.2. Equipo de trabajo", level=2)
    titulo_tabla(doc, "Tabla 6. Roles, habilidades y responsabilidades.")
    tabla(doc, ["Rol", "Integrante", "Habilidades requeridas", "Responsabilidades"],
          C.EQUIPO, anchos=[2.8, 2.6, 5.2, 5.4], tamano=9)
    parrafo(doc, C.NOTA_EQUIPO, cursiva=True, tamano=10)

    doc.add_heading("3.3. Herramientas de gestión", level=2)
    parrafo(doc,
            "El seguimiento del trabajo se llevó en Jira Software, en un "
            "proyecto de tipo Scrum con su tablero, su Product Backlog y un "
            "Sprint Backlog por iteración. Las diez historias y las cuatro "
            "épicas están cargadas allí con sus criterios de aceptación y sus "
            "puntos de historia.")
    parrafo(doc,
            "Se eligió Jira por ser el estándar en la industria y por "
            "distinguir de forma nativa entre épicas, historias y tareas, algo "
            "que un tablero genérico de tarjetas no hace. El repositorio "
            "incluye además el archivo historias_usuario_jira.csv, que permite "
            "reconstruir el tablero completo mediante la importación de datos "
            "de Jira.")
    titulo_tabla(doc, "Tabla 7. Herramientas de soporte al proyecto.")
    tabla(doc, ["Herramienta", "Uso en el proyecto"],
          [["Jira Software", "Product Backlog, Sprint Backlog, tablero y seguimiento de las historias."],
           ["GitHub", "Alojamiento del repositorio e historial de cambios."],
           ["GitHub Actions", "Integración continua: ejecuta las 99 pruebas ante cada cambio."],
           ["Visual Studio Code", "Entorno de desarrollo."],
           ["Microsoft Teams", "Ceremonias de Scrum y comunicación."],
           ["OneDrive institucional", "Entrega del documento y del video."]],
          anchos=[4.2, 11.4], tamano=9.5)

    doc.add_heading("3.4. Épicas", level=2)
    parrafo(doc,
            "Las historias se agruparon en cuatro épicas según el área "
            "funcional que abordan. La épica del circuito de préstamos es la "
            "más grande porque concentra las reglas de negocio del sistema.")
    titulo_tabla(doc, "Tabla 8. Épicas del Release 1.")
    tabla(doc, ["ID", "Épica", "Descripción", "Historias", "Puntos"],
          C.EPICAS, anchos=[1.3, 3.4, 6.6, 2.9, 1.4], tamano=9)

    doc.add_heading("3.5. Ceremonias de Scrum", level=2)
    parrafo(doc,
            "Las ceremonias se planificaron con fecha y hora fijas para los "
            "tres sprints: Sprint 1 del 22 de junio al 3 de julio, Sprint 2 del "
            "6 al 17 de julio y Sprint 3 del 20 al 31 de julio de 2026.")
    titulo_tabla(doc, "Tabla 9. Calendario de ceremonias.")
    tabla(doc, ["Ceremonia", "Cuándo", "Duración", "Horario", "Participantes",
                "Propósito"],
          C.CEREMONIAS, anchos=[2.4, 4.0, 1.6, 1.9, 2.3, 3.4], tamano=8.5)

    doc.add_heading("3.6. Historias de usuario", level=2)
    parrafo(doc,
            "Las diez historias siguen el formato «Como… quiero… para…», que "
            "obliga a expresar quién necesita la función y para qué, no solo "
            "qué hace el sistema. Cada una lleva sus criterios de aceptación y "
            "su estimación en puntos de historia según la sucesión de "
            "Fibonacci (3, 5, 8), donde el valor refleja la complejidad y no "
            "las horas de trabajo.")
    parrafo(doc,
            "HU-04 es la historia de mayor puntuación (8) porque concentra las "
            "cuatro validaciones del negocio; HU-01, HU-03, HU-07 y HU-10 son "
            "las más sencillas (3). El total asciende a 45 puntos.")

    for hid, titulo, narrativa, criterios, puntos, epica, sprint, prioridad in C.HISTORIAS:
        doc.add_heading(f"{hid} — {titulo}", level=3)
        parrafo(doc, narrativa, cursiva=True, espacio_despues=4)
        tabla(doc, ["Épica", "Sprint", "Prioridad", "Puntos de historia"],
              [[epica, sprint, prioridad, str(puntos)]],
              anchos=[3.9, 3.9, 3.9, 3.9], tamano=9.5)
        parrafo(doc, "Criterios de aceptación:", negrita=True, tamano=10.5,
                espacio_despues=2)
        for i, criterio in enumerate(criterios, 1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_after = Pt(2)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(f"{i}. {criterio}")
            r.font.size = Pt(10.5)

    parrafo(doc, "", espacio_despues=6)
    titulo_tabla(doc, "Tabla 10. Resumen de historias y distribución por sprint.")
    resumen = [[h[0], h[1], h[5], h[6], str(h[4])] for h in C.HISTORIAS]
    resumen.append(["", "TOTAL", "", "", str(sum(h[4] for h in C.HISTORIAS))])
    tabla(doc, ["ID", "Historia", "Épica", "Sprint", "Puntos"], resumen,
          anchos=[1.4, 6.6, 2.4, 2.4, 1.5], tamano=9)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def plan_pruebas(doc):
    doc.add_heading("4. Plan de pruebas", level=1)
    parrafo(doc,
            "Este plan define qué se prueba, cómo se decide si el resultado es "
            "aceptable, con qué herramientas, quién lo ejecuta y cuándo. Su "
            "objetivo no es demostrar que el sistema funciona, sino intentar "
            "que falle antes de que lo haga en producción.")

    doc.add_heading("4.1. Requerimientos y matriz de trazabilidad", level=2)
    parrafo(doc,
            "Los requerimientos funcionales y no funcionales se listaron en el "
            "apartado 2.7. La matriz siguiente los relaciona con la historia "
            "que los origina y con las pruebas que los verifican, de modo que "
            "ningún requerimiento quede sin comprobar.")
    titulo_tabla(doc, "Tabla 11. Matriz de trazabilidad requerimiento – historia – prueba.")
    matriz = [
        ["RF-01", "HU-01", "test_devuelve_el_catalogo_inicial, test_lista_el_catalogo", "API + E2E"],
        ["RF-02", "HU-02", "TestCrearLibro (7 casos), test_registrar_un_libro", "API + E2E"],
        ["RF-03", "HU-03", "TestListarLibros (5 casos), test_buscar_por_titulo", "API + E2E"],
        ["RF-04", "HU-04", "TestCalcularVencimiento (7 casos)", "Unitaria"],
        ["RF-05", "HU-04", "TestValidarPrestamo (8), TestRegistrarPrestamo (11)", "Unitaria + API + E2E"],
        ["RF-06", "HU-05, HU-06", "TestCalcularMora (7), TestDevolucion (6), TestMoraPorRetraso", "Unitaria + API"],
        ["RF-07", "HU-07", "TestConsultarPrestamos (3 casos)", "API"],
        ["RF-08", "HU-09", "TestSocios (6), test_registrar_un_socio", "API + E2E"],
        ["RF-09", "HU-10", "TestAutenticacion (3), TestAcceso (4)", "API + E2E"],
        ["RF-10", "HU-08", "TestReportes (4), test_muestra_los_indicadores", "API + E2E"],
    ]
    tabla(doc, ["Requerimiento", "Historia", "Pruebas que lo verifican", "Nivel"],
          matriz, anchos=[2.4, 2.4, 7.6, 3.6], tamano=9)

    doc.add_heading("4.2. Criterios de aceptación y de rechazo", level=2)
    parrafo(doc, "Criterios de entrada (cuándo se puede empezar a probar)",
            negrita=True, espacio_despues=2)
    vinetas(doc, C.CRITERIOS_ENTRADA)
    parrafo(doc, "Criterios de aceptación (cuándo se aprueba el Release)",
            negrita=True, espacio_despues=2)
    vinetas(doc, C.CRITERIOS_ACEPTACION)
    parrafo(doc, "Criterios de rechazo (cuándo se devuelve a desarrollo)",
            negrita=True, espacio_despues=2)
    vinetas(doc, C.CRITERIOS_RECHAZO)
    parrafo(doc, "Criterios de suspensión (cuándo se detienen las pruebas)",
            negrita=True, espacio_despues=2)
    vinetas(doc, C.CRITERIOS_SUSPENSION)

    doc.add_heading("4.3. Herramientas de pruebas y su justificación", level=2)
    titulo_tabla(doc, "Tabla 12. Herramientas de pruebas y motivo de su elección.")
    tabla(doc, ["Herramienta", "Se usa para", "Justificación de la elección"],
          C.HERRAMIENTAS_PRUEBAS, anchos=[3.4, 3.2, 9.4], tamano=9)

    doc.add_heading("4.4. Cronograma de ejecución de pruebas", level=2)
    titulo_tabla(doc, "Tabla 13. Cronograma de pruebas manuales y automatizadas.")
    tabla(doc, ["Actividad", "Inicio", "Fin", "Tipo", "Responsable"],
          C.CRONOGRAMA_PRUEBAS, anchos=[6.4, 2.1, 2.1, 2.4, 3.0], tamano=9)

    doc.add_heading("4.5. Plantilla de casos de prueba", level=2)
    parrafo(doc,
            "Se adoptó la plantilla siguiente, basada en la norma IEEE 829, "
            "para documentar cada caso de forma uniforme. Los campos son los "
            "mínimos que permiten a otra persona reproducir la prueba sin "
            "preguntar nada.")
    titulo_tabla(doc, "Tabla 14. Plantilla estándar para documentar un caso de prueba.")
    tabla(doc, ["Campo", "Descripción del campo"],
          [["ID del caso", "Identificador único, formato CP-000."],
           ["Historia asociada", "Historia de usuario que origina el caso."],
           ["Título", "Qué se comprueba, en una frase."],
           ["Precondición", "Estado del sistema y datos necesarios antes de empezar."],
           ["Pasos", "Acciones numeradas, en el orden exacto de ejecución."],
           ["Resultado esperado", "Qué debe ocurrir, en términos observables."],
           ["Tipo", "Unitaria, API, E2E; manual o automatizada."],
           ["Resultado obtenido", "Superado o fallido, con la fecha de ejecución."]],
          anchos=[4.0, 11.6], tamano=9.5)

    parrafo(doc,
            "A continuación se documentan ocho casos representativos con la "
            "plantilla ya rellena. Se eligieron los que cubren las reglas de "
            "negocio más delicadas y los casos negativos, que son los que "
            "realmente ponen a prueba el sistema.", espacio_despues=8)

    for cid, hu, titulo, precondicion, pasos, esperado, tipo, resultado in C.CASOS_PRUEBA:
        titulo_tabla(doc, f"{cid} — {titulo}")
        pasos_texto = "\n".join(f"{i}. {p}" for i, p in enumerate(pasos, 1))
        tabla(doc, ["Campo", "Contenido"],
              [["ID del caso", cid],
               ["Historia asociada", hu],
               ["Título", titulo],
               ["Precondición", precondicion],
               ["Pasos", pasos_texto],
               ["Resultado esperado", esperado],
               ["Tipo", tipo],
               ["Resultado obtenido", f"{resultado} (15/08/2026)"]],
              anchos=[3.6, 12.0], tamano=9)

    doc.add_heading("4.6. Equipo de pruebas y responsabilidades", level=2)
    titulo_tabla(doc, "Tabla 15. Responsabilidades en el proceso de pruebas.")
    tabla(doc, ["Rol", "Responsable", "Responsabilidades en las pruebas"],
          C.EQUIPO_PRUEBAS, anchos=[3.4, 3.0, 9.6], tamano=9)

    doc.add_heading("4.7. Plan de automatización de pruebas", level=2)
    parrafo(doc,
            "La automatización se organizó en tres niveles siguiendo la "
            "pirámide de pruebas descrita por Fowler (2012): cuanto más abajo "
            "está una prueba, más rápida y estable es, y más de ellas conviene "
            "tener.")
    titulo_tabla(doc, "Tabla 16. Niveles de automatización implementados.")
    tabla(doc, ["Nivel", "Cantidad", "Qué cubre", "Herramienta", "Velocidad",
                "Detalle"],
          C.PLAN_AUTOMATIZACION,
          anchos=[2.6, 1.8, 3.2, 2.4, 1.9, 3.7], tamano=8.5)
    parrafo(doc, "Estrategia y decisiones tomadas", negrita=True, espacio_despues=2)
    vinetas(doc, C.ESTRATEGIA_AUTOMATIZACION)

    doc.add_heading("4.8. Ejecución de las pruebas y evidencias", level=2)
    parrafo(doc,
            "La suite completa se ejecutó el 15 de agosto de 2026 sobre el "
            "Release 1. Los resultados son los siguientes.")
    titulo_tabla(doc, "Tabla 17. Resultados de la última ejecución completa.")
    tabla(doc, ["Nivel", "Total", "Superadas", "Fallidas", "Porcentaje"],
          C.RESULTADOS, anchos=[4.6, 2.6, 2.8, 2.6, 3.0], tamano=9.5)

    parrafo(doc, "Cobertura de código", negrita=True, espacio_despues=2)
    titulo_tabla(doc, "Tabla 18. Cobertura por módulo.")
    tabla(doc, ["Módulo", "Responsabilidad", "Cobertura"],
          C.COBERTURA, anchos=[5.2, 6.8, 3.6], tamano=9.5)
    parrafo(doc, C.NOTA_COBERTURA, cursiva=True, tamano=10)

    parrafo(doc, "Defectos detectados por la automatización", negrita=True,
            espacio_despues=2)
    parrafo(doc,
            "Los dos defectos siguientes fueron encontrados por las pruebas "
            "automatizadas durante el desarrollo, no por un usuario. Se "
            "documentan porque son la evidencia concreta de que el plan de "
            "pruebas cumplió su función.")
    titulo_tabla(doc, "Tabla 19. Defectos detectados, causa y corrección.")
    tabla(doc, ["ID", "Severidad", "Descripción", "Causa raíz", "Corrección",
                "Estado"],
          C.DEFECTOS, anchos=[1.4, 1.8, 3.9, 3.9, 3.2, 1.4], tamano=8.5)

    parrafo(doc, "Evidencias disponibles en el repositorio", negrita=True,
            espacio_despues=2)
    vinetas(doc, [
        "evidencias/demo/demo_biblioitla.mp4 — video del recorrido completo por "
        "el Release 1, grabado automáticamente con Playwright.",
        "evidencias/capturas/ — una captura de pantalla por cada prueba E2E, "
        "generada automáticamente al finalizar cada caso.",
        "evidencias/videos/ — grabación en video de cada prueba E2E ejecutándose.",
        "evidencias/cobertura/index.html — informe de cobertura navegable, línea "
        "por línea.",
        "Pestaña Actions del repositorio — historial de todas las ejecuciones de "
        "la suite en integración continua.",
    ])
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def entregables(doc):
    doc.add_heading("5. Demostración y entregables", level=1)
    parrafo(doc,
            "El video de demostración recorre el incremento del Release 1 en el "
            "orden en que trabaja el bibliotecario: inicio de sesión, panel de "
            "indicadores, catálogo, alta de un libro, búsqueda, registro de un "
            "socio, préstamo, rechazo de un préstamo que incumple una regla, "
            "devolución y actualización de los indicadores.")
    parrafo(doc,
            "Se incluye a propósito el intento de préstamo rechazado: mostrar "
            "solo el camino correcto no demuestra que las reglas de negocio "
            "estén implementadas.")
    titulo_tabla(doc, "Tabla 20. Enlaces de la entrega.")
    tabla(doc, ["Entregable", "Enlace"],
          [["Repositorio del código fuente", C.REPOSITORIO],
           ["Tablero de Jira con las historias de usuario", C.TABLERO],
           ["Código de las pruebas automatizadas",
            f"{C.REPOSITORIO}/tree/main/pruebas"],
           ["Integración continua (ejecuciones de la suite)",
            f"{C.REPOSITORIO}/actions"],
           ["Video de demostración del Release 1", C.VIDEO]],
          anchos=[5.6, 10.0], tamano=9.5)
    parrafo(doc,
            "Los enlaces se colocan además en la opción «Texto en línea» de la "
            "plataforma, como exige el reglamento de entrega.", cursiva=True,
            tamano=10)


def conclusiones(doc):
    doc.add_heading("6. Conclusiones", level=1)
    for texto in C.CONCLUSIONES:
        parrafo(doc, texto)


def bibliografia(doc):
    doc.add_heading("7. Bibliografía", level=1)
    for referencia in C.BIBLIOGRAFIA:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(-1.0)   # sangría francesa
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(referencia)
        r.font.size = Pt(11)


# ===========================================================================
def main():
    doc = Document()
    preparar_estilos(doc)

    seccion = doc.sections[0]
    seccion.top_margin = Cm(2.5)
    seccion.bottom_margin = Cm(2.5)
    seccion.left_margin = Cm(2.5)
    seccion.right_margin = Cm(2.5)
    numerar_paginas(seccion)

    portada(doc)
    indice(doc)
    introduccion(doc)
    estrategia(doc)
    scrum(doc)
    plan_pruebas(doc)
    entregables(doc)
    conclusiones(doc)
    bibliografia(doc)

    doc.save(SALIDA)
    print(f"Documento generado: {SALIDA}")
    print(f"Tamano: {SALIDA.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    main()
